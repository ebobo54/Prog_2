from pyforms.basewidget import BaseWidget
from pyforms.controls import ControlButton, ControlTextArea
from electron import Electron
from neitron import Neitron
from proton import Proton
from docx import Document

class Calculator(BaseWidget):
    def __init__(self):
        super().__init__('Калькулятор свойств частиц')

        self.electron_button = ControlButton('Рассчитать свойства электрона', default=self.calculate_electron)
        self.neitron_button = ControlButton('Рассчитать свойства нейтрона', default=self.calculate_neitron)
        self.proton_button = ControlButton('Рассчитать свойства протона', default=self.calculate_proton)
        self._output_area = ControlTextArea('Результаты')
        self.formset = [('electron_button', 'neitron_button', 'proton_button', ), '_output_area']

    def calculate_electron(self):
        particle = Electron()
        specific_charge = particle.specific_charge()
        compton = particle.compton()
        result_text = f"Удельный заряд электрона: {specific_charge:.2e} Кл/кг\nКомптоновская длина: {compton:.2e} м"
        self._output_area.value = result_text
        doc = Document()
        doc.add_paragraph(result_text)
        doc.save('електрон.doc')

    def calculate_neitron(self):
        particle = Neitron()
        specific_charge = particle.specific_charge()
        compton = particle.compton()
        result_text = f"Удельный заряд электрона: {specific_charge:.2e} Кл/кг\nКомптоновская длина: {compton:.2e} м"
        self._output_area.value = result_text
        doc = Document()
        doc.add_paragraph(result_text)
        doc.save('нейтрон.doc')

    def calculate_proton(self):
        particle = Proton()
        specific_charge = particle.specific_charge()
        compton = particle.compton()
        result_text = f"Удельный заряд электрона: {specific_charge:.2e} Кл/кг\nКомптоновская длина: {compton:.2e} м"
        self._output_area.value = result_text
        doc = Document()
        doc.add_paragraph(result_text)
        doc.save('протон.doc')

if __name__ == "__main__":
    from pyforms import start_app
    start_app(Calculator)

